"""Create the reference.docx template for pandoc."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# --- Page margins: top=2cm, bottom=1cm, left=2cm, right=1cm ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1)

    # --- Page number in footer, bottom center ---
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = run2._r.makeelement(qn("w:instrText"), {})
    instrText.text = " PAGE "
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run3._r.append(fldChar2)

# --- Set default font to Roboto ---
style = doc.styles["Normal"]
font = style.font
font.name = "Roboto"
font.size = Pt(11)

# --- Heading styles: Roboto Bold ---
for i in range(1, 5):
    heading = doc.styles[f"Heading {i}"]
    heading.font.name = "Roboto"

# --- Code style: Consolas ---
code_style = doc.styles.add_style("Source Code", 1)
code_font = code_style.font
code_font.name = "Consolas"
code_font.size = Pt(9)

# Inline code character style
char_style = doc.styles.add_style("Verbatim Char", 2)
char_font = char_style.font
char_font.name = "Consolas"
char_font.size = Pt(9)

# Dummy paragraph so template is valid
doc.add_paragraph("Template document - styles only", style="Normal")

doc.save("template/reference.docx")
print("Template saved: template/reference.docx")
